"""
File processing utilities for extracting text from PDF and DOCX files
"""
import os
from PyPDF2 import PdfReader
from docx import Document
import pdfplumber


def process_uploaded_file(filepath):
    """
    Extract text content from uploaded file
    
    Args:
        filepath (str): Path to the uploaded file
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If file type is not supported
    """
    file_extension = os.path.splitext(filepath)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(filepath)
    elif file_extension in ['.docx', '.doc']:
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def extract_text_from_pdf(filepath):
    """
    Extract text from PDF file using multiple methods for better accuracy
    
    Args:
        filepath (str): Path to PDF file
        
    Returns:
        str: Extracted text content
    """
    text = ""
    
    # Try pdfplumber first (better for tables and complex layouts)
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed: {e}, trying PyPDF2...")
        
        # Fallback to PyPDF2
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 also failed: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    return text.strip()


def extract_text_from_docx(filepath):
    """
    Extract text from DOCX file
    
    Args:
        filepath (str): Path to DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        doc = Document(filepath)
        text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        return "\n".join(text)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {e}")


def validate_file_type(filename, allowed_extensions=None):
    """
    Validate if file type is allowed
    
    Args:
        filename (str): Name of the file
        allowed_extensions (set): Set of allowed extensions
        
    Returns:
        bool: True if file type is allowed
    """
    if allowed_extensions is None:
        allowed_extensions = {'pdf', 'docx', 'doc'}
    
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions
